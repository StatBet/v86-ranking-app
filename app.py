import streamlit as st
import pandas as pd
import docx
from datetime import datetime
from supabase import create_client
from skrall_badge_engine import apply_skrall_badges
from v8x_postprocess import apply_v8x_postprocess
from environment_live_engine import environment_label
from badge_engine import assign_badges, calculate_spike_score, get_round_spikes
from loser_badge_helpers import apply_loser_badges_to_race
from debug_live_lopp_sums import get_live_lopp_sum_debug
#from loppbadge_sum_helpers import get_sum_loppbadge
from badge_rules import (
    get_race_metrics,
    get_loppbadge,
    apply_model_probabilities,
)

from live_rank_probability_engine_v2 import (
    assign_rank_probabilities,
)


from scripts.ranking_engine_v3 import (
    parse_input,
    add_dynamic_scores,
    calculate_total_score,
    scoring_rules
)

from scripts.parser_atg_new import parse_new_atg_format


st.set_page_config(page_title="V86 Ranking App", layout="wide")
st.title("V86 Ranking App")

SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

supabase = create_client(
    SUPABASE_URL,
    SUPABASE_KEY,
)

try:
    supabase.storage.from_("startfiler").list(
        "",
        {
            "limit": 1,
            "offset": 0,
        },
    )
    
except Exception as error:
    st.error(f"Supabase-anslutning misslyckades: {error}")
    st.stop()


def read_docx(file):
    doc = docx.Document(file)
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                if cell.text.strip():
                    cells.append(cell.text.strip())
            if cells:
                text.append(" | ".join(cells))

    return "\n".join(text)


def clean_atg_header(raw_data):
    if "Avdelning 1," in raw_data:
        raw_data = raw_data.split("Avdelning 1,", 1)[1]
        raw_data = "Avdelning 1," + raw_data
    elif "Avdelning 1" in raw_data:
        raw_data = raw_data.split("Avdelning 1", 1)[1]
        raw_data = "Avdelning 1" + raw_data

    return raw_data


def int_slider(label, value, min_value=-50, max_value=50, key=None):
    return st.slider(
        label,
        min_value=min_value,
        max_value=max_value,
        value=int(value),
        step=1,
        key=key
    )


st.sidebar.title("Poängpanel")
st.sidebar.caption("Ändra poäng här. Resultatet uppdateras automatiskt.")


with st.sidebar.expander("Spel%", expanded=True):
    use_spel_percent = st.toggle(
        "Använd spel% i totalpoäng",
        value=False,
        key="sidebar_use_spel_percent"
    )

    for i, value in enumerate(scoring_rules["spel_percent_points"]):
        scoring_rules["spel_percent_points"][i] = int_slider(
            f"Spel% grupp {i + 1}",
            value,
            0,
            40,
            f"sidebar_spel_{i}"
        )

    scoring_rules["spel_percent_group_threshold"] = int_slider(
        "Gruppgräns %-enheter",
        scoring_rules["spel_percent_group_threshold"],
        0,
        10,
        "sidebar_spel_threshold"
    )


with st.sidebar.expander("Snittodds"):
    for i, row in enumerate(scoring_rules["avg_odds_ranges"]):
        row["points"] = int_slider(
            f"Odds {row['min']}–{row['max']}",
            row["points"],
            0,
            40,
            f"sidebar_odds_{i}"
        )


with st.sidebar.expander("Seger%"):
    for i, value in enumerate(scoring_rules["win_percent_points"]):
        scoring_rules["win_percent_points"][i] = int_slider(
            f"Seger% grupp {i + 1}",
            value,
            0,
            40,
            f"sidebar_win_{i}"
        )

    scoring_rules["win_percent_group_threshold"] = int_slider(
        "Gruppgräns %-enheter",
        scoring_rules["win_percent_group_threshold"],
        0,
        10,
        "sidebar_win_threshold"
    )


with st.sidebar.expander("Plats%"):
    for i, value in enumerate(scoring_rules["place_percent_points"]):
        scoring_rules["place_percent_points"][i] = int_slider(
            f"Plats% grupp {i + 1}",
            value,
            0,
            40,
            f"sidebar_place_{i}"
        )

    scoring_rules["place_percent_group_threshold"] = int_slider(
        "Gruppgräns %-enheter",
        scoring_rules["place_percent_group_threshold"],
        0,
        10,
        "sidebar_place_threshold"
    )


with st.sidebar.expander("Prissumma"):
    for i, value in enumerate(scoring_rules["prize_money_points"]):
        scoring_rules["prize_money_points"][i] = int_slider(
            f"Prissumma grupp {i + 1}",
            value,
            0,
            40,
            f"sidebar_prize_{i}"
        )

    scoring_rules["prize_money_group_threshold_percent"] = int_slider(
        "Gruppgräns %",
        scoring_rules["prize_money_group_threshold_percent"],
        0,
        20,
        "sidebar_prize_threshold"
    )


with st.sidebar.expander("Senaste 5 prispengar"):
    for i, row in enumerate(scoring_rules.get("recent_prize_ranges", [])):
        row["points"] = int_slider(
            f"Min {row['min']} kr",
            row["points"],
            0,
            30,
            f"sidebar_recent_prize_{i}"
        )


with st.sidebar.expander("Kuskmodell"):
    st.caption("Automatisk kuskpoäng hämtas från config/driver_stats.json.")

    scoring_rules["driver_min_starts"] = int_slider(
        "Minsta antal lopp för kuskpoäng",
        scoring_rules.get("driver_min_starts", 70),
        0,
        500,
        "sidebar_driver_min_starts"
    )

    scoring_rules["driver_mid_starts"] = int_slider(
        "Gräns nivå 2",
        scoring_rules.get("driver_mid_starts", 150),
        0,
        1000,
        "sidebar_driver_mid_starts"
    )

    scoring_rules["driver_high_starts"] = int_slider(
        "Gräns nivå 3",
        scoring_rules.get("driver_high_starts", 300),
        0,
        1500,
        "sidebar_driver_high_starts"
    )

    scoring_rules["driver_low_multiplier"] = st.slider(
        "Multiplier låg nivå",
        min_value=0.0,
        max_value=2.0,
        value=float(scoring_rules.get("driver_low_multiplier", 0.75)),
        step=0.05,
        key="sidebar_driver_low_multiplier"
    )

    scoring_rules["driver_mid_multiplier"] = st.slider(
        "Multiplier mellan nivå",
        min_value=0.0,
        max_value=2.0,
        value=float(scoring_rules.get("driver_mid_multiplier", 1.0)),
        step=0.05,
        key="sidebar_driver_mid_multiplier"
    )

    scoring_rules["driver_high_multiplier"] = st.slider(
        "Multiplier hög nivå",
        min_value=0.0,
        max_value=2.0,
        value=float(scoring_rules.get("driver_high_multiplier", 1.25)),
        step=0.05,
        key="sidebar_driver_high_multiplier"
    )


with st.sidebar.expander("Form"):
    for placement in list(scoring_rules["form_points"].keys()):
        scoring_rules["form_points"][placement] = int_slider(
            f"Placering {placement}",
            scoring_rules["form_points"][placement],
            0,
            25,
            f"sidebar_form_{placement}"
        )


with st.sidebar.expander("Senaste start"):
    for placement in list(scoring_rules["latest_start_points"].keys()):
        scoring_rules["latest_start_points"][placement] = int_slider(
            f"Senaste start {placement}",
            scoring_rules["latest_start_points"][placement],
            0,
            25,
            f"sidebar_latest_{placement}"
        )


with st.sidebar.expander("Starter"):
    for i, row in enumerate(scoring_rules["starts_points"]):
        row["points"] = int_slider(
            f"{row['min']}–{row['max']} starter",
            row["points"],
            -30,
            30,
            f"sidebar_starts_{i}"
        )


with st.sidebar.expander("Rekord"):
    st.caption("Rekordpoängen beräknas nu relativt inom loppet.")


    
with st.sidebar.expander("Vagn / Skor / Manuell"):
    scoring_rules["american_wagon_bonus"] = int_slider(
        "Amerikansk vagn bonus",
        scoring_rules["american_wagon_bonus"],
        0,
        30,
        "sidebar_wagon_bonus"
    )

    scoring_rules["american_wagon_max_recent_count"] = int_slider(
        "Max amerikansk senaste 5",
        scoring_rules["american_wagon_max_recent_count"],
        0,
        5,
        "sidebar_wagon_recent"
    )

    manual_shoe_bonus = int_slider(
        "Skorpoäng per ikryssad häst",
        0,
        -20,
        30,
        "sidebar_manual_shoe_bonus"
    )

    manual_stallform_bonus = int_slider(
        "Stallformpoäng per ikryssad häst",
        8,
        0,
        30,
        "sidebar_manual_stallform_bonus"
    )


with st.sidebar.expander("Inaktivitet"):
    inactivity_days_limit = int_slider(
        "Dagar utan start",
        90,
        0,
        365,
        "sidebar_inactivity_days"
    )

    inactivity_penalty = int_slider(
        "Poängavdrag",
        -5,
        -50,
        0,
        "sidebar_inactivity_penalty"
    )


with st.sidebar.expander("Galopp / Kön"):
    scoring_rules["gallop_penalty"] = int_slider(
        "Galoppavdrag",
        scoring_rules["gallop_penalty"],
        -30,
        0,
        "sidebar_gallop_penalty"
    )

    scoring_rules["gallop_penalty_min_count"] = int_slider(
        "Min antal galopper",
        scoring_rules["gallop_penalty_min_count"],
        1,
        5,
        "sidebar_gallop_min"
    )

    scoring_rules["gender_penalty_sto_mixed"] = int_slider(
        "Sto mot hingst/vallack",
        scoring_rules["gender_penalty_sto_mixed"],
        -30,
        0,
        "sidebar_gender_penalty"
    )


with st.sidebar.expander("Tillägg distans"):
    scoring_rules["distance_addition_penalty"]["1640"] = int_slider(
        "1640m per 20m",
        scoring_rules["distance_addition_penalty"]["1640"],
        -40,
        0,
        "sidebar_dist_1640"
    )

    scoring_rules["distance_addition_penalty"]["2140"] = int_slider(
        "2140m per 20m",
        scoring_rules["distance_addition_penalty"]["2140"],
        -40,
        0,
        "sidebar_dist_2140"
    )

    scoring_rules["distance_addition_penalty"]["2640"] = int_slider(
        "2640m+ per 20m",
        scoring_rules["distance_addition_penalty"]["2640"],
        -40,
        0,
        "sidebar_dist_2640"
    )

st.subheader("📚 Startfiler")

try:
    files = supabase.storage.from_("startfiler").list(
        "",
        {
            "limit": 1000,
            "offset": 0,
        },
    )

    file_names = sorted(
        [
            file["name"]
            for file in files
            if file.get("name", "").lower().endswith(".txt")
        ],
        reverse=True,
    )

except Exception as error:
    st.error(f"Kunde inte läsa startfiler:\n\n{error}")
    file_names = []


selected_library_file = st.selectbox(
    "Välj en sparad startfil",
    options=["— Välj fil —"] + file_names,
    key="supabase_startfile_select",
)


uploaded_file = st.file_uploader(
    "Ladda upp startlista (.txt eller .docx)",
    type=["txt", "docx"],
    key="main_file_uploader",
)


raw_data = None
source_name = None


# En lokalt uppladdad fil har företräde.
if uploaded_file is not None:

    source_name = uploaded_file.name

    if uploaded_file.name.lower().endswith(".txt"):
        file_bytes = uploaded_file.getvalue()

        try:
            raw_data = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw_data = file_bytes.decode("latin-1")

        raw_data = clean_atg_header(raw_data)

    else:
        raw_data = read_docx(uploaded_file)
        raw_data = clean_atg_header(raw_data)


# Om ingen lokal fil är uppladdad kan en fil från biblioteket användas.
elif selected_library_file != "— Välj fil —":

    source_name = selected_library_file

    try:
        file_bytes = supabase.storage.from_("startfiler").download(
            selected_library_file
        )

        try:
            raw_data = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw_data = file_bytes.decode("latin-1")

        raw_data = clean_atg_header(raw_data)

    except Exception as error:
        st.error(f"Kunde inte hämta startfilen:\n\n{error}")


if raw_data is not None:

    races = parse_input(raw_data)

    races = [
        r for r in races
        if r["race"].get("track") != "UNKNOWN"
        and r["race"].get("distance", 0) > 0
        and r["race"].get("start") != "unknown"
    ]

    if not races or any(len(r["horses"]) <= 1 for r in races):
        races = parse_new_atg_format(raw_data)

    races = [
        r for r in races
        if r["race"].get("track") != "UNKNOWN"
        and r["race"].get("distance", 0) > 0
        and r["race"].get("start") != "unknown"
    ]


    # Spara automatiskt en giltig lokalt uppladdad TXT-fil i Supabase.
    if (
        uploaded_file is not None
        and uploaded_file.name.lower().endswith(".txt")
        and races
        and uploaded_file.name not in file_names
    ):
        try:
            supabase.storage.from_("startfiler").upload(
                path=uploaded_file.name,
                file=uploaded_file.getvalue(),
                file_options={
                    "content-type": "text/plain",
                },
            )

            st.success(
                f"☁️ {uploaded_file.name} sparades automatiskt "
                "i det gemensamma biblioteket."
            )

        except Exception as error:
            st.warning(
                "Rankingen kan fortsätta, men filen kunde inte sparas "
                f"i biblioteket:\n\n{error}"
            )


    all_spike_candidates = []

    st.write("Antal tecken inläst:", len(raw_data))
    st.write("Antal avdelningar hittade:", len(races))

    processed_races = []
    summary_placeholder = st.empty()

    live_debug_rows = []

    for race_data in races:
        race = race_data["race"]
        horses = race_data["horses"]

        horses = add_dynamic_scores(horses, race)

        horses = apply_loser_badges_to_race(horses)

        st.subheader(
            f"{race['track']} - Avdelning {race['race_no']} - "
            f"{race['distance']}m ({race['start']})"
        )

       

        #if horses:
            #st.write("DEBUG första häst keys:", list(horses[0].keys()))
            #st.write("DEBUG första häst:", horses[0])

        #sum_badge = get_sum_loppbadge(horses)

        #if sum_badge and sum_badge.get("badge"):
            #st.info(
                #f"{sum_badge['badge']} | TotalSum: {sum_badge['total_sum']} | SpikeSum: {sum_badge['spike_sum']}"
            #)

        #sum_badge = get_sum_loppbadge(horses)

        #if sum_badge and sum_badge.get("badge"):
            #st.info(
                #f"{sum_badge['badge']} | TotalSum: {sum_badge['total_sum']} | SpikeSum: {sum_badge['spike_sum']}"
            #)

        with st.expander(f"Manuell skorjustering - Avdelning {race['race_no']}"):
            cols = st.columns(3)

            for idx, horse in enumerate(horses):
                number = horse.get("number", 0)
                name = horse.get("horse", "")

                checked = cols[idx % 3].checkbox(
                    f"{number} {name}",
                    value=False,
                    key=f"shoe_checkbox_{race['race_no']}_{idx}_{name}"
                )

                horse["shoe_score"] = manual_shoe_bonus if checked else 0

        with st.expander(f"Manuell stallform - Avdelning {race['race_no']}"):
            cols = st.columns(3)

            for idx, horse in enumerate(horses):
                number = horse.get("number", 0)
                name = horse.get("horse", "")

                checked = cols[idx % 3].checkbox(
                    f"{number} {name}",
                    value=False,
                    key=f"stall_checkbox_{race['race_no']}_{idx}_{name}"
                )

                horse["stallform_score"] = manual_stallform_bonus if checked else 0

        for horse in horses:
            history = horse.get("history", [])

            if history:
                latest_date = history[0].get("date", "")

                try:
                    latest_date_obj = datetime.strptime(latest_date, "%Y-%m-%d")
                    days_since = (datetime.today() - latest_date_obj).days

                    if days_since > inactivity_days_limit:
                        horse["inactivity_score"] = inactivity_penalty
                    else:
                        horse["inactivity_score"] = 0

                except Exception:
                    horse["inactivity_score"] = 0
            else:
                horse["inactivity_score"] = 0

        if not use_spel_percent:
            for horse in horses:
                horse["spel_score"] = 0

        for horse in horses:
            horse["total_score"] = calculate_total_score(horse)

        race_for_badges = dict(race)
        race_for_badges["horses"] = horses

        horses = assign_badges(horses, race_for_badges)

        for h in horses:
            h["race_no"] = race.get("race_no", "")
            h["race_track"] = race.get("track", "")
            h["spike_score"] = calculate_spike_score(h, race_for_badges)

        horses = apply_loser_badges_to_race(horses)

        debug_sums = get_live_lopp_sum_debug(horses)

        is_compact = (
            debug_sums["total_sum"] <= 1165
            or debug_sums["spike_sum"] <= 1097
        )

        race["is_compact"] = is_compact
        race["debug_sums"] = debug_sums  

        #st.caption(
            #f"DEBUG lopp-summor | "        
            #f"TotalSum: {debug_sums['total_sum']} | "
            #f"SpikeSum: {debug_sums['spike_sum']} | "
            #f"CompactTotal: {debug_sums['compact_total']} | "
            #f"CompactSpike: {debug_sums['compact_spike']} | "
            #f"ChaosTotal: {debug_sums['chaos_total']} | "
            #f"ChaosSpike: {debug_sums['chaos_spike']}"
        #)

        live_debug_rows.append({
            "race_no": race["race_no"],
            "track": race["track"],
            "total_sum": debug_sums["total_sum"],
            "spike_sum": debug_sums["spike_sum"],
        })

        pd.DataFrame(live_debug_rows).to_csv("live_lopp_sums_debug.csv", index=False)

        race_output_placeholder = st.empty()

        processed_races.append({
            "race": race,
            "horses": horses,
            "placeholder": race_output_placeholder
        })

    # Hybrid först: skapar leaf/environment-data som Skrällmotorn behöver.
    top_spikes = get_round_spikes(processed_races)

    # Fryst Skrällmotor: Spår 1 + Spår 2 bygger kandidatpoolen
    # med korrekt live leaf/environment-data.
    processed_races = apply_skrall_badges(processed_races)

    # V8X fryst slutlager:
    # Startpoäng -> EPS -> Spike/EPS-rescues -> final Skräll/Main+Rescue.
    # total_score/spike_score/hybrid ändras inte.
    processed_races = apply_v8x_postprocess(processed_races)
    # Endast visning – påverkar inte hybridmotorn eller spikvalen.
    hybrid_audit_rows = []

    for race_data in processed_races:
        race = race_data["race"]

        ranked_horses = sorted(
            race_data["horses"],
            key=lambda horse: horse.get("total_score", 0),
            reverse=True,
        )

        if not ranked_horses:
            continue

        rank1 = ranked_horses[0]

        metrics = get_race_metrics(ranked_horses)
        loppbadge = get_loppbadge(metrics)

        spike_percent = rank1.get(
            "hybrid_spike_percent"
        )

        environment_percent = rank1.get(
            "hybrid_environment_percent"
        )

        available_percentages = [
            float(value)
            for value in [
                spike_percent,
                environment_percent,
            ]
            if value is not None
        ]

        best_percent = (
            max(available_percentages)
            if available_percentages
            else None
        )

        selected_position = rank1.get(
            "hybrid_spike_position"
        )

        hybrid_audit_rows.append({
            "Avd": race.get("race_no", ""),
            "Rank 1": rank1.get(
                "horse",
                rank1.get("name", ""),
            ),
            "Lopptyp": loppbadge.get(
                "label",
                "Okänd",
            ),
            "Spread 1–8": round(
                float(metrics.get("spread_1_8", 0)),
                1,
            ),
            "Score gap 1–2": round(
                float(metrics.get("gap_1_2", 0)),
                1,
            ),
            "Spike %": (
                round(float(spike_percent), 2)
                if spike_percent is not None
                else None
            ),
            "Miljö %": (
                round(float(environment_percent), 2)
                if environment_percent is not None
                else None
            ),
            "Bästa motor": (
                "ENVIRONMENT"
                if (
                    environment_percent is not None
                    and (
                        spike_percent is None
                        or float(environment_percent)
                        > float(spike_percent)
                    )
                )
                else "SPIKE"
            ),
            "Bästa %": (
                round(best_percent, 2)
                if best_percent is not None
                else None
            ),

            "Spike match": (
                "JA"
                if rank1.get(
                    "hybrid_spike_profile_matched",
                    False,
                )
                else "NEJ"
            ),
            "Spikeprofil": rank1.get(
                "hybrid_spike_profile",
                "",
            ),
            
            "Vald": (
                "JA"
                if selected_position is not None
                else "NEJ"
            ),
            "Spikplats": selected_position or "",
        })

    SHOW_HYBRID_DEBUG = False

    if SHOW_HYBRID_DEBUG:
        with st.expander(
            "🔍 Hybridkontroll – alla lopp",
            expanded=True,
        ):
            hybrid_audit_df = pd.DataFrame(
                hybrid_audit_rows
            )

            if not hybrid_audit_df.empty:
                hybrid_audit_df = (
                    hybrid_audit_df
                    .sort_values(
                        ["Bästa %", "Avd"],
                        ascending=[False, True],
                    )
                    .reset_index(drop=True)
                )

                st.dataframe(
                    hybrid_audit_df,
                    width="stretch",
                    hide_index=True,
                )


    with summary_placeholder.container():
        st.subheader("🎯 Omgångens spikförslag")

        for i, horse in enumerate(top_spikes, start=1):
            if i > 3:
                continue

            badge = "🟩 Toppspik" if i <= 2 else "🟦 Spik"

            selected_engine = horse.get(
                "hybrid_selected_engine",
                "UNKNOWN"
            )

            selected_percent = horse.get(
                "hybrid_selected_percent"
            )

            if selected_percent is not None:
                chance_text = (
                    f"**Hybridchans: {float(selected_percent):.1f}%** "
                    f"({selected_engine})"
                )
            else:
                chance_text = "**Hybridchans: saknas**"

            st.caption(
                f"Hybridmotor: {selected_engine} | "
                f"Spikeprofil: "
                f"{horse.get('hybrid_spike_percent', 'saknas')} | "
                f"Rank1-miljö: "
                f"{horse.get('hybrid_environment_percent', 'saknas')}"
            )

            value_names = horse.get("value_candidate_names", [])
            value_text = ""

            if value_names:
                value_text = (
                    "\n\nSystem-only valuekandidat(er): "
                    + ", ".join(value_names)
                )

            st.markdown(
                f"""
    {badge} **{horse.get("horse", horse.get("name", ""))}**
    — Avd {horse.get("race_no", "")}
    — Rank: {horse.get("_model_rank_live", horse.get("model_rank", ""))}
    — Score: {round(horse.get("total_score", 0), 1)}
    — SpikeScore: {round(horse.get("spike_score", 0), 1)}
    — Spel%: {horse.get("percent", 0)}%
    {chance_text}
    {value_text}
    """
            )

    for race_data in processed_races:
        # Legacy badges, race environment and probability engine keep using the
        # original total_score order they were trained/validated on.
        horses = sorted(
            race_data["horses"],
            key=lambda x: x.get("total_score", 0),
            reverse=True
        )

        from rank68_badge_helpers import apply_rank68_badges

        for h in horses:
            h = apply_rank68_badges(h)
            
        horses = apply_loser_badges_to_race(horses)

        metrics = get_race_metrics(horses)
        loppbadge = get_loppbadge(metrics)

        rank1_horse = horses[0] if horses else {}

        leaf_id = rank1_horse.get(
            "hybrid_environment_leaf_id"
        )

        try:
            leaf_id = int(leaf_id)
        except (TypeError, ValueError):
            leaf_id = None

        environment_badges = {
            20: {
                "type": "success",
                "text": "🟢 Favoritmiljö | Rank 1–3: 76%",
            },
            19: {
                "type": "success",
                "text": "🟢 Favoritmiljö | Rank 1–3: 73%",
            },
            8: {
                "type": "info",
                "text": "🔵 Kompakt miljö | Rank 1–5: 80%",
            },
            4: {
                "type": "warning",
                "text": "🟠 Jämnt lopp | Rank 6+: 33%",
            },
            10: {
                "type": "warning",
                "text": "🟠 Skrällmiljö | Rank 6+: 31%",
            },
            5: {
                "type": "error",
                "text": "🔴 Kaosmiljö | Rank 6+: 43%",
            },
        }

        environment_badge = environment_badges.get(
            leaf_id
        )

        is_compact = race_data["race"].get("is_compact", False)
        debug_sums = race_data["race"].get("debug_sums", {})

        is_open = loppbadge.get("label") == "Öppet lopp"
        is_3horse = loppbadge.get("label") == "3-hästarslopp"

        if is_open:
            primary_race_type = "OPEN"
        elif is_compact:
            primary_race_type = "COMPACT"
        elif is_3horse:
            primary_race_type = "3HORSE"
        else:
            primary_race_type = "STANDARD"

        if primary_race_type == "OPEN":
            probability_badge = {
                "label": "Öppet lopp"
            }

        elif primary_race_type == "COMPACT":
            probability_badge = {
                "label": "Kompakt lopp"
            }

        elif primary_race_type == "3HORSE":
            probability_badge = {
                "label": "3-hästarslopp"
            }

        else:
            probability_badge = {
                "label": "Öppet lopp"
            }    

        rank1_horse = horses[0] if horses else {}

        leaf_id = rank1_horse.get(
            "hybrid_environment_leaf_id",
            rank1_horse.get(
                "rank1_environment_leaf_id"
            ),
        )

        spike_percent = rank1_horse.get(
            "hybrid_spike_percent"
        )

        environment_percent = rank1_horse.get(
            "hybrid_environment_percent"
        )

        available_percentages = [
            float(value)
            for value in [
                spike_percent,
                environment_percent,
            ]
            if value is not None
        ]

        hybrid_rank1_percent = (
            max(available_percentages)
            if available_percentages
            else None
        )

        probability_result = assign_rank_probabilities(
            horses=horses,
            leaf_id=leaf_id,
            hybrid_rank1_percent=hybrid_rank1_percent,
        )

        horses = probability_result["horses"]

        # Only now switch to the V8X final rank for user-facing ranking/output.
        horses = sorted(
            horses,
            key=lambda h: h.get("_final_rank", h.get("base_model_rank", 999)),
        )
        race_data["horses"] = horses

        for idx, h in enumerate(horses, start=1):
            h["display_rank"] = idx

        rows = []

        for idx, h in enumerate(horses, start=1):
            rows.append({
                "Rank": idx,
                "Nr": h.get("number", 0),
                "Häst": h.get("horse", ""),

                "Badges": "  ".join(
                    b for b in h.get("badges", [])
                    if "Top5" not in b
                    and "Topp 5" not in b
                    and "Topp5" not in b
                    and "TOP5" not in b
                ),

                "Tot": h.get("total_score", 0),

                "SpikeScore": round(
                    h.get("spike_score", 0),
                    1
                ),

                "Form": h.get("form_score", 0),

                "SP-rank": h.get(
                    "start_points_rank",
                    ""
                ),

                "EPS-rank": h.get(
                    "eps_rank",
                    ""
                ),

                "Speed": h.get("speed_score", 0),
                "AvgTid": h.get("avg_time", ""),
                "Stallform": h.get("stallform_score", 0),
                "Senaste": h.get("latest_start_score", 0),
                "Spårpoäng": h.get("post_score", 0),
                "Kusk": h.get("driver_score", 0),
                "Kuskbyte": h.get("driver_change_score", 0),
                "Rek": h.get("record_score", 0),
                "Starter": h.get("starts_score", 0),
                "Seger%": h.get("win_score", 0),
                "Plats%": h.get("place_score", 0),
                "Spel%": h.get("spel_score", 0),
                "Pris": h.get("prize_money_score", 0),
                "Senaste pris": h.get("recent_prize_score", 0),
                "Klass": h.get("class_change_score", 0),
                "Odds": h.get("avg_odds_score", 0),
                "SnittOdds": h.get("avg_odds", ""),
                "Vagn": h.get("wagon_score", 0),
                "Skor": h.get("shoe_score", 0),
                "Inaktiv": h.get("inactivity_score", 0),
                "Manuell": h.get("custom_score", 0),
                "Tillägg": h.get("distance_addition_score", 0),
                "Kön": h.get("gender_score", 0),
                "Galopp": h.get("gallop_score", 0),
                "Prissumma": h.get("prize_money", 0),
                "St": h.get("starts", 0),
                "V%": h.get("win_percent", 0),
                "P%": h.get("place_percent", 0),
                "Spelad %": h.get("percent", 0),
                "Vagn idag": h.get("equipment", ""),
                "Kusk namn": h.get("driver", "")
            })

        df = pd.DataFrame(rows)

        with race_data["placeholder"].container():

            # =================================================
            # ENVIRONMENT V2 - SLUTLIG LIVE-MILJÖ
            #
            # VIKTIGT:
            # Detta är ENDAST Environment V2-visningen.
            # Hybrid V3 / gamla miljöleafs / spikuttagning
            # lämnas helt orörda.
            # =================================================

            env_race = race_data.get(
                "race",
                {},
            )

            environment_v2 = env_race.get(
                "final_environment_v2",
                "UT",
            )

            environment_profile_v2 = env_race.get(
                "final_environment_profile_v2",
            )

            initial_environment_v2 = env_race.get(
                "initial_environment_v2",
                environment_v2,
            )

            initial_leaf_v2 = env_race.get(
                "initial_environment_leaf_v2",
            )

            final_leaf_v2 = env_race.get(
                "final_environment_leaf_v2",
            )

            environment_text_v2 = environment_label(
                environment_v2,
                environment_profile_v2,
            )

            # Miljöerna visas tydligt men påverkar INTE
            # Hybrid V3:s spikuttagning.
            if environment_v2 == "Favoritrank":
                st.success(
                    environment_text_v2
                )

            elif environment_v2 == "Solid":
                st.success(
                    environment_text_v2
                )

            elif environment_v2 == "Neutral":
                st.info(
                    environment_text_v2
                )

            elif environment_v2 == "Öppet":
                st.warning(
                    environment_text_v2
                )

            elif environment_v2 == "Kaos":
                st.warning(
                    environment_text_v2
                )

            elif environment_v2 == "Extrem kaos":
                st.error(
                    environment_text_v2
                )

            else:
                st.caption(
                    environment_text_v2
                )

            # Visa migration endast när själva miljön ändrats.
            if (
                initial_environment_v2
                != environment_v2
            ):
                st.caption(
                    "Environment V2: "
                    f"{initial_environment_v2} "
                    f"→ {environment_v2} "
                    f"| leaf "
                    f"{initial_leaf_v2} "
                    f"→ {final_leaf_v2}"
                )


            
            st.dataframe(
                df,
                width="stretch",
                height=(len(df) + 1) * 35 + 3,
                hide_index=True,
                column_config={
                    "Rank": st.column_config.NumberColumn(
                        "Rank",
                        pinned=True
                    ),
                    "Nr": st.column_config.NumberColumn(
                        "Nr",
                        pinned=True
                    ),
                    "Häst": st.column_config.TextColumn(
                        "Häst",
                        pinned=True
                    ),
                }
            )
