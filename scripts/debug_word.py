from docx import Document

file_path = r"C:\Users\Grinvald\Desktop\Ranking v85\data\raw_data.docx"

doc = Document(file_path)

print("START DEBUG\n")

print("PARAGRAPHS COUNT:", len(doc.paragraphs))
print("TABLES COUNT:", len(doc.tables))

print("\n--- FIRST 20 PARAGRAPHS ---\n")

for i, p in enumerate(doc.paragraphs[:20]):
    print(i, repr(p.text))