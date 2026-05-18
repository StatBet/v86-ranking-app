from docx import Document
import re

FILE = r"data/raw_data.docx"

doc = Document(FILE)

print("\n================ FULL RAW DUMP ================\n")

all_text = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        all_text.append(text)
        print(f"P{i:03}: {text}")

print("\n================ TABLE DUMP ================\n")

for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti} ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"{ti}.{ri}: {cells}")

print("\n================ HUNT HORSENAME PATTERN ================\n")

# väldigt enkel häst-detektion (för debug)
horse_candidates = []

pattern = re.compile(r"[A-Za-zÅÄÖåäö].*\d$")

for t in all_text:
    if pattern.search(t):
        horse_candidates.append(t)

print(f"Horse candidates found: {len(horse_candidates)}\n")

for h in horse_candidates[:50]:
    print(h)