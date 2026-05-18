from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parent.parent
WORD_FILE = BASE / "data" / "raw_data.docx"

doc = Document(WORD_FILE)

print("\n================ PARAGRAPHS =================\n")

for i, p in enumerate(doc.paragraphs[:200]):
    if p.text.strip():
        print(f"{i}: {p.text}")

print("\n================ TABLES =================\n")

for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti} ---\n")

    for ri, row in enumerate(table.rows):
        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if row_text:
            print(f"{ti}.{ri}: {row_text}")