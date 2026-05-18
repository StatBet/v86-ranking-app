from docx import Document

file_path = r"C:\Users\Grinvald\Desktop\Ranking v85\data\raw_data.docx"

doc = Document(file_path)

print("\n================ WORD DEBUG TEST ================\n")

print("Paragraphs:", len(doc.paragraphs))
print("Tables:", len(doc.tables))

print("\n================ PARAGRAPHS (RAW) ================\n")

for i, p in enumerate(doc.paragraphs[:30]):
    print(f"{i:02d}:", repr(p.text))

print("\n================ TABLES (RAW) ================\n")

for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti + 1} ---")

    for ri, row in enumerate(table.rows):
        row_text = [cell.text.strip() for cell in row.cells]
        print(f"{ti}.{ri}:", " | ".join(row_text))